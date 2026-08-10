"""
标准查重 API
"""

import io
import re
from datetime import datetime
from typing import Dict, List, Optional
from urllib.parse import quote

from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from loguru import logger
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from pydantic import BaseModel, Field

from app.schemas.base import Fail, Success, SuccessExtra

router = APIRouter(prefix="/deduplication", tags=["标准查重"])


# ──────────────────────────────────────────────
# Pydantic Schemas
# ──────────────────────────────────────────────

class SimilarityTags(BaseModel):
    """相似度标签"""
    同系列标准: bool = Field(False, description="是否属于同一系列标准")
    标准化对象一致: bool = Field(False, description="标准化的对象/主题是否一致")
    通用专用关系: bool = Field(False, description="一个是通用标准，另一个是针对特定领域/产品的专用标准")
    适用范围重叠: bool = Field(False, description="两个标准在实际应用中是否可能被同一个主体同时采用或实施（有实质性的用户群体交集和应用场景重叠）")


class SimilarStandard(BaseModel):
    """相似标准（基于标签评判）"""
    id: str
    cname: str
    use_range: Optional[str]
    standard_no: Optional[str]
    vector_score: float = Field(description="向量相似度（仅供参考）")
    tags: SimilarityTags = Field(description="相似度标签")
    relation_desc: str = Field(description="关系说明")
    has_sim_cache: bool = Field(default=False, description="是否已有全文相似度缓存")
    has_ai_comparison_cache: bool = Field(default=False, description="是否已有AI比对缓存")
    cache_time: Optional[str] = Field(None, description="缓存时间")
    has_res: bool = Field(default=False, description="是否有标准原文")


def _make_excel_header_style():
    fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    font = Font(color="FFFFFF", bold=True, size=12)
    alignment = Alignment(horizontal="center", vertical="center")
    border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )
    return fill, font, alignment, border


def _apply_header_style(ws_row, fill, font, alignment, border):
    for cell in ws_row:
        cell.fill = fill
        cell.font = font
        cell.alignment = alignment
        cell.border = border


_EXCEL_ILLEGAL_CHARS_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")


def _clean_excel_text(value):
    """openpyxl 不允许 0x00-0x08 / 0x0b / 0x0c / 0x0e-0x1f 这些控制字符。
    数据里偶尔会夹这些不可见字符，写入前剥掉。"""
    if value is None or not isinstance(value, str):
        return value
    return _EXCEL_ILLEGAL_CHARS_RE.sub("", value)


def _generate_batch_excel(record, detail_rows: list, similar_rows: list) -> io.BytesIO:
    """根据批次记录生成 Excel 工作簿，返回 BytesIO。

    detail_rows: 主表行（list[dict]，键见下方 SQL）
    similar_rows: 相似明细拍平行（list[dict]，键见下方 SQL）

    采用 openpyxl write_only 模式：cell 一次性 stream 到文件而不驻留对象树，
    30 万行规模的内存占用从 ~2GB 降到 ~50MB。
    """
    wb = Workbook(write_only=True)
    fill, font, alignment, border = _make_excel_header_style()
    left_wrap_align = Alignment(horizontal="left", vertical="center", wrap_text=True)

    # ── 1. 批次汇总 ──
    from openpyxl.cell import WriteOnlyCell
    ws_summary = wb.create_sheet(title="批次汇总")

    def _hcell(ws, value):
        c = WriteOnlyCell(ws, value=value)
        c.fill = fill
        c.font = font
        c.alignment = alignment
        c.border = border
        return c

    def _dcell(ws, value, align=None, bd=None):
        c = WriteOnlyCell(ws, value=_clean_excel_text(value))
        if align is not None:
            c.alignment = align
        if bd is not None:
            c.border = bd
        return c

    ws_summary.append([
        _hcell(ws_summary, h)
        for h in ["批次ID", "创建时间", "总数", "成功数", "失败数", "高相似数", "状态"]
    ])
    ws_summary.append([
        record.id,
        record.create_time.strftime("%Y-%m-%d %H:%M:%S") if record.create_time else "",
        record.total_count,
        record.success_count,
        record.failed_count,
        record.duplicate_count,
        record.status,
    ])
    for col, width in zip('ABCDEFG', [12, 20, 10, 10, 10, 12, 12]):
        ws_summary.column_dimensions[col].width = width

    # ── 2. 详细结果 ──
    ws_detail = wb.create_sheet(title="详细结果")
    detail_headers = ["序号", "标准编号", "标准名称", "适用范围", "状态",
                      "相似标准数", "需要关注", "总体评价", "建议措施", "错误信息"]
    ws_detail.append([_hcell(ws_detail, h) for h in detail_headers])

    for idx, r in enumerate(detail_rows, 1):
        found = bool(r.get("found"))
        need_attention = bool(r.get("need_attention"))
        status = "未找到" if not found else ("需要关注" if need_attention else "无需关注")
        similar_count = r.get("similar_count") if found else ""
        ws_detail.append([
            _dcell(ws_detail, idx, left_wrap_align, border),
            _dcell(ws_detail, r.get("standard_no") or "", left_wrap_align, border),
            _dcell(ws_detail, r.get("standard_name") or "", left_wrap_align, border),
            _dcell(ws_detail, r.get("use_range") or "", left_wrap_align, border),
            _dcell(ws_detail, status, left_wrap_align, border),
            _dcell(ws_detail, similar_count, left_wrap_align, border),
            _dcell(ws_detail, "是" if need_attention else "否", left_wrap_align, border),
            _dcell(ws_detail, r.get("general_evaluation") or "", left_wrap_align, border),
            _dcell(ws_detail, r.get("suggestion") or "", left_wrap_align, border),
            _dcell(ws_detail, r.get("error") or "", left_wrap_align, border),
        ])
    for col, width in zip('ABCDEFGHIJ', [8, 20, 40, 40, 12, 12, 12, 50, 50, 30]):
        ws_detail.column_dimensions[col].width = width

    # ── 3. 相似标准明细 ──
    ws_similar = wb.create_sheet(title="相似标准明细")
    similar_headers = [
        "原标准编号", "原标准名称", "原标准适用范围", "需要关注",
        "相似标准编号", "相似标准名称", "相似标准适用范围",
        "LLM评分", "同系列标准", "标准化对象一致",
        "通用专用关系", "适用范围重叠", "关系说明",
    ]
    ws_similar.append([_hcell(ws_similar, h) for h in similar_headers])

    for r in similar_rows:
        ws_similar.append([
            _dcell(ws_similar, r.get("main_standard_no") or "", left_wrap_align, border),
            _dcell(ws_similar, r.get("main_standard_name") or "", left_wrap_align, border),
            _dcell(ws_similar, r.get("main_use_range") or "", left_wrap_align, border),
            _dcell(ws_similar, "是" if r.get("main_need_attention") else "否", left_wrap_align, border),
            _dcell(ws_similar, r.get("sn") or "", left_wrap_align, border),
            _dcell(ws_similar, r.get("cn") or "", left_wrap_align, border),
            _dcell(ws_similar, r.get("ur") or "", left_wrap_align, border),
            _dcell(ws_similar, round(float(r.get("ls") or 0), 4), left_wrap_align, border),
            _dcell(ws_similar, "是" if r.get("t1") else "否", left_wrap_align, border),
            _dcell(ws_similar, "是" if r.get("t2") else "否", left_wrap_align, border),
            _dcell(ws_similar, "是" if r.get("t3") else "否", left_wrap_align, border),
            _dcell(ws_similar, "是" if r.get("t4") else "否", left_wrap_align, border),
            _dcell(ws_similar, r.get("rd") or "", left_wrap_align, border),
        ])
    for col, width in zip('ABCDEFGHIJKLM', [20, 35, 40, 12, 20, 35, 40, 12, 15, 15, 15, 15, 50]):
        ws_similar.column_dimensions[col].width = width

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output


# ──────────────────────────────────────────────
# Word 生成 — 辅助函数（模块级）
# ──────────────────────────────────────────────

def _set_font_yahei(element):
    """设置元素字体为微软雅黑"""
    from docx.oxml.ns import qn
    if hasattr(element, 'font'):
        element.font.name = '微软雅黑'
        element.font._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


def _safe_iter_runs(element):
    if element is None or not hasattr(element, 'runs') or element.runs is None:
        return []
    return element.runs


def _safe_iter_paragraphs(element):
    if element is None or not hasattr(element, 'paragraphs') or element.paragraphs is None:
        return []
    return element.paragraphs


def _safe_iter_cells(element):
    if element is None or not hasattr(element, 'cells') or element.cells is None:
        return []
    return element.cells


def _set_cell_font(cell):
    """设置表格单元格字体为微软雅黑"""
    for paragraph in _safe_iter_paragraphs(cell):
        for run in _safe_iter_runs(paragraph):
            _set_font_yahei(run)


def _add_indicator_table(doc, indicators: List[Dict], cols: List[str], title: str):
    """通用：在 Word 文档中添加指标表格"""
    if not indicators:
        return

    doc.add_paragraph()
    heading = doc.add_heading(title, 4)
    for run in _safe_iter_runs(heading):
        _set_font_yahei(run)

    table = doc.add_table(rows=len(indicators) + 1, cols=len(cols))
    table.style = 'Light Grid Accent 1'

    for col_idx, col_name in enumerate(cols):
        table.cell(0, col_idx).text = col_name
    for cell in _safe_iter_cells(table.rows[0]):
        for para in _safe_iter_paragraphs(cell):
            for run in _safe_iter_runs(para):
                run.bold = True
                _set_font_yahei(run)

    for row_idx, indicator in enumerate(indicators, 1):
        if isinstance(indicator, dict):
            for col_idx, key in enumerate(cols):
                cell_keys = {
                    '序号': str(row_idx),
                    '指标大类': indicator.get('category') or '',
                    '指标名称': indicator.get('name') or '',
                    '源标准要求': indicator.get('source_requirement') or '',
                    '目标标准要求': indicator.get('target_requirement') or '',
                    '差异分析': indicator.get('change_analysis') or '',
                    '要求': indicator.get('requirement') or '',
                }
                table.cell(row_idx, col_idx).text = cell_keys.get(key, '')

    for row in table.rows:
        for cell in _safe_iter_cells(row):
            _set_cell_font(cell)


async def _build_word_document(standard_result: Dict, standard_no: str) -> io.BytesIO:
    """根据查重结果生成 Word 报告文档，返回 BytesIO"""
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.shared import RGBColor

    from app.models.standard import StandardCacheAI, StandardCacheSim

    doc = Document()

    # 设置全局字体
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    for i in range(10):
        try:
            s = doc.styles[f'Heading {i}']
            s.font.name = '微软雅黑'
            s._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        except Exception:
            pass

    # ── 标题 ──
    title = doc.add_heading('标准相似度分析报告', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in _safe_iter_runs(title):
        _set_font_yahei(run)

    # ── 一、标准基本信息 ──
    h1 = doc.add_heading('一、标准基本信息', 1)
    for run in _safe_iter_runs(h1):
        _set_font_yahei(run)

    table_basic = doc.add_table(rows=3, cols=2)
    table_basic.style = 'Light Grid Accent 1'
    table_basic.cell(0, 0).text = '标准编号'
    table_basic.cell(0, 1).text = standard_result.get('standard_no') or ''
    table_basic.cell(1, 0).text = '标准名称'
    table_basic.cell(1, 1).text = standard_result.get('standard_name') or ''
    table_basic.cell(2, 0).text = '适用范围'
    table_basic.cell(2, 1).text = standard_result.get('use_range') or ''
    for row in table_basic.rows:
        for cell in _safe_iter_cells(row):
            _set_cell_font(cell)

    # ── 二、总体评价 ──
    h2 = doc.add_heading('二、总体评价', 1)
    for run in _safe_iter_runs(h2):
        _set_font_yahei(run)

    need_attention = standard_result.get('need_attention', False)
    p_attention = doc.add_paragraph()
    run_label = p_attention.add_run('是否需要关注：')
    run_label.bold = True
    _set_font_yahei(run_label)
    run_val = p_attention.add_run('需要关注' if need_attention else '无需关注')
    run_val.font.color.rgb = RGBColor(255, 0, 0) if need_attention else RGBColor(0, 128, 0)
    run_val.bold = True
    _set_font_yahei(run_val)

    overall_eval = standard_result.get('general_evaluation', '')
    if overall_eval:
        p_eval = doc.add_paragraph(overall_eval)
        for run in _safe_iter_runs(p_eval):
            _set_font_yahei(run)

    # ── 三、建议措施 ──
    suggestions = standard_result.get('suggestion', '')
    if suggestions:
        h3 = doc.add_heading('三、建议措施', 1)
        for run in _safe_iter_runs(h3):
            _set_font_yahei(run)
        p_sugg = doc.add_paragraph(suggestions)
        for run in _safe_iter_runs(p_sugg):
            _set_font_yahei(run)

    # ── 四、相似标准详情 ──
    similar_standards = standard_result.get('similar_standards', [])
    if similar_standards:
        h4 = doc.add_heading('四、相似标准详情', 1)
        for run in _safe_iter_runs(h4):
            _set_font_yahei(run)
        p_count = doc.add_paragraph(f'共发现 {len(similar_standards)} 个相似标准')
        for run in _safe_iter_runs(p_count):
            _set_font_yahei(run)

        for idx, similar in enumerate(similar_standards, 1):
            h_sim = doc.add_heading(f'{idx}. {similar.get("standard_no", "未知编号")}', 2)
            for run in _safe_iter_runs(h_sim):
                _set_font_yahei(run)

            # 相似标准基本信息
            table_sim = doc.add_table(rows=4, cols=2)
            table_sim.style = 'Light Grid Accent 1'
            table_sim.cell(0, 0).text = '标准名称'
            table_sim.cell(0, 1).text = similar.get('cname') or ''
            table_sim.cell(1, 0).text = '适用范围'
            table_sim.cell(1, 1).text = similar.get('use_range') or ''
            table_sim.cell(2, 0).text = '向量相似度'
            table_sim.cell(2, 1).text = f'{similar.get("vector_score", 0):.2%}'
            tags = similar.get('tags', {})
            active_tags = [k for k, v in tags.items() if v]
            table_sim.cell(3, 0).text = '疑似交叉类型'
            table_sim.cell(3, 1).text = '、'.join(active_tags) if active_tags else '无'
            for row in table_sim.rows:
                for cell in _safe_iter_cells(row):
                    _set_cell_font(cell)

            relation_desc = similar.get('relation_desc', '')
            if relation_desc:
                p_rel = doc.add_paragraph(f'关系说明：{relation_desc}')
                for run in _safe_iter_runs(p_rel):
                    _set_font_yahei(run)

            # 全文相似度缓存 & AI比对缓存
            target_std_no = similar.get('standard_no')
            if target_std_no:
                full_text_cache = await StandardCacheSim.filter(
                    source_standard_no=standard_no,
                    target_standard_no=target_std_no,
                    is_valid=True
                ).first()

                if full_text_cache:
                    h_ft = doc.add_heading('全文相似度分析', 3)
                    for run in _safe_iter_runs(h_ft):
                        run.font.color.rgb = RGBColor(0, 112, 192)
                        _set_font_yahei(run)

                    table_ft = doc.add_table(rows=4, cols=2)
                    table_ft.style = 'Light Grid Accent 1'
                    table_ft.cell(0, 0).text = '相似度'
                    cell_sim_val = table_ft.cell(0, 1)
                    cell_sim_val.text = f'{full_text_cache.similarity_percentage:.2f}%'
                    for para in _safe_iter_paragraphs(cell_sim_val):
                        for run in _safe_iter_runs(para):
                            run.bold = True
                            run.font.color.rgb = (
                                RGBColor(255, 0, 0)
                                if full_text_cache.similarity_percentage > 50
                                else RGBColor(0, 0, 255)
                            )
                            _set_font_yahei(run)
                    table_ft.cell(1, 0).text = '匹配句子数'
                    table_ft.cell(1, 1).text = str(full_text_cache.matched_sentence_count)
                    table_ft.cell(2, 0).text = '源标准总句子数'
                    table_ft.cell(2, 1).text = str(full_text_cache.source_total_sentence_count)
                    table_ft.cell(3, 0).text = '目标标准总句子数'
                    table_ft.cell(3, 1).text = str(full_text_cache.target_total_sentence_count)
                    for row in table_ft.rows:
                        for cell in _safe_iter_cells(row):
                            _set_cell_font(cell)
                    doc.add_paragraph()

                ai_cache = await StandardCacheAI.filter(
                    source_standard_no=standard_no,
                    target_standard_no=target_std_no,
                    is_valid=True
                ).first()

                if ai_cache:
                    h_ai = doc.add_heading('AI智能比对分析', 3)
                    for run in _safe_iter_runs(h_ai):
                        run.font.color.rgb = RGBColor(0, 176, 80)
                        _set_font_yahei(run)

                    table_ai = doc.add_table(rows=1, cols=2)
                    table_ai.style = 'Light Grid Accent 1'
                    table_ai.cell(0, 0).text = '关系结论'
                    table_ai.cell(0, 1).text = ai_cache.relationship or '无'
                    for row in table_ai.rows:
                        for cell in _safe_iter_cells(row):
                            _set_cell_font(cell)

                    _add_indicator_table(
                        doc, ai_cache.matched_indicators or [],
                        ['序号', '指标大类', '指标名称', '源标准要求', '目标标准要求', '差异分析'],
                        '匹配的指标对比'
                    )
                    _add_indicator_table(
                        doc, ai_cache.source_unique_indicators or [],
                        ['序号', '指标大类', '指标名称', '要求'],
                        '源标准独有指标'
                    )
                    _add_indicator_table(
                        doc, ai_cache.target_unique_indicators or [],
                        ['序号', '指标大类', '指标名称', '要求'],
                        '目标标准独有指标'
                    )

            doc.add_paragraph()

    # 页脚
    doc.add_paragraph()
    footer = doc.add_paragraph()
    run_footer = footer.add_run(f'生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    _set_font_yahei(run_footer)
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output


# ──────────────────────────────────────────────
# 路由处理函数
# ──────────────────────────────────────────────

@router.get("/batch-summary", summary="批量相似度分析汇总统计")
async def get_batch_summary():
    """获取批量相似度分析全局 KPI 统计"""
    from datetime import datetime, timedelta

    from app.models.standard import StandardDuplicateBatch

    try:
        total_batches = await StandardDuplicateBatch.all().count()
        seven_days_ago = datetime.now() - timedelta(days=7)
        recent_7d_batches = await StandardDuplicateBatch.filter(create_time__gte=seven_days_ago).count()

        all_records = await StandardDuplicateBatch.all().values("total_count", "duplicate_count")
        total_duplicate = sum(r["duplicate_count"] for r in all_records)
        total_analyzed = sum(r["total_count"] for r in all_records)

        return Success(data={
            "total_batches": total_batches,
            "recent_7d_batches": recent_7d_batches,
            "total_duplicate": total_duplicate,
            "total_analyzed": total_analyzed,
        }, msg="获取汇总统计成功")
    except Exception as e:
        return Fail(code="5000", msg=f"获取汇总统计失败: {str(e)}")


@router.get("/batch-history", summary="批量相似度分析历史记录")
async def get_batch_history(
        current: int = 1,
        size: int = 10,
        batch_name: Optional[str] = None,
        status: Optional[str] = None,
        pool_id: Optional[int] = None,
        start_time: Optional[str] = None,
        end_time: Optional[str] = None,
):
    """获取批量相似度分析历史记录列表"""
    from datetime import datetime

    from app.models.standard import StandardDuplicateBatch, StandardDuplicateName, StandardPoolCheck

    try:
        qs = StandardDuplicateBatch.all()
        if batch_name:
            qs = qs.filter(batch_name__icontains=batch_name)
        if status:
            qs = qs.filter(status=status)
        if pool_id is not None:
            qs = qs.filter(pool_id=pool_id)
        if start_time:
            qs = qs.filter(create_time__gte=datetime.fromisoformat(start_time))
        if end_time:
            qs = qs.filter(create_time__lte=datetime.fromisoformat(end_time))

        offset = (current - 1) * size
        total = await qs.count()
        records = await qs.offset(offset).limit(size).order_by("-create_time")

        # 批量查询涉及的 pool_id
        pool_ids = {r.pool_id for r in records if r.pool_id and r.pool_id > 0}
        pool_name_map: dict[int, str] = {}
        if pool_ids:
            pools = await StandardPoolCheck.filter(id__in=pool_ids).values("id", "pool_name")
            pool_name_map = {p["id"]: p["pool_name"] for p in pools}

        def get_pool_name(pool_id: int | None) -> str:
            if pool_id is None:
                return "-"
            if pool_id == -1:
                return "全库"
            return pool_name_map.get(pool_id, f"池#{pool_id}")

        # 对处于 processing 的批次，实时统计 task_status 分布（用于前端进度条）
        processing_ids = [r.id for r in records if r.status == "processing"]
        progress_map: dict[int, dict] = {}
        if processing_ids:
            name_rows = await StandardDuplicateName.filter(
                batch_id__in=processing_ids
            ).values("batch_id", "task_status")
            for row in name_rows:
                bid = row["batch_id"]
                ts = row["task_status"] or "pending"
                bucket = progress_map.setdefault(
                    bid, {"pending": 0, "running": 0, "done": 0, "failed": 0}
                )
                bucket[ts] = bucket.get(ts, 0) + 1

        records_data = [
            {
                "id": r.id,
                "batch_name": r.batch_name,
                "create_time": r.create_time.isoformat() if r.create_time else None,
                "update_time": r.update_time.isoformat() if r.update_time else None,
                "total_count": r.total_count,
                "success_count": r.success_count,
                "failed_count": r.failed_count,
                "duplicate_count": r.duplicate_count,
                "status": r.status,
                "pool_id": r.pool_id,
                "pool_name": get_pool_name(r.pool_id),
                "remark": r.remark,
                "progress": progress_map.get(r.id),
            }
            for r in records
        ]
        return SuccessExtra(data={"records": records_data}, total=total, current=current, size=size, msg="获取历史记录成功")
    except Exception as e:
        return Fail(code="5000", msg=f"获取历史记录失败: {str(e)}")


# tags 字段白名单（防止 SQL 注入和未知 tag）
_VALID_TAGS = {
    "同系列标准", "标准化对象一致", "通用专用关系",
    "适用范围重叠", "技术领域相同", "标准类型相似",
}

# 跨域筛选映射：(源标准领域, 目标标准领域)
_CROSS_DOMAIN: dict[str, tuple[str, str]] = {
    "国家标准对行业标准": ("国家标准", "行业标准"),
    "行业标准对国家标准": ("行业标准", "国家标准"),
}


def _escape_like(kw: str) -> str:
    """转义 LIKE 中的特殊字符 % _ \\，避免用户输入被解释为通配符。"""
    return kw.replace("\\", "\\\\").replace("%", "\\%").replace("_", "\\_")


@router.get("/batch-detail/{batch_id}", summary="批量相似度分析详情")
async def get_batch_detail(
    batch_id: int,
    current: int = 1,
    size: int = 10,
    only_need_attention: bool = False,
    only_found: bool = True,
    filter_tags: str = None,
    hidden_tags: str = None,
    filter_tags_mode: str = 'any',
    sort_field: str = None,
    sort_order: str = None,
    filter_standard_no: str = None,
    filter_standard_name: str = None,
    has_res_filter: str = None,
    main_has_res_filter: str = None,
    filter_std_domain: str = None,
):
    """
    获取批量相似度分析详情（支持分页、筛选、排序）。

    实现要点：过滤 / 排序 / 分页全部下推 SQL，避免拉全量。tag 筛选用
    JSON_TABLE 拍平 + EXISTS 表达；similar_count 排序用 JSON_LENGTH。
    has_res_filter 仅过滤 similar_standards 数组内元素，不剔除主记录
    （保持旧版本语义）。main_has_res_filter='has_res' 含义为
    "主标准在 jgh_pdf 中且至少一个相似标准在 jgh_pdf 中"，需在拿到
    当前页后做后置过滤（依赖 JSON 内 standard_no 与 jgh_pdf 的连接，
    含义复杂故不下推 SQL）。
    """
    from tortoise import Tortoise

    from app.models.standard import (
        StandardBaseInfo,
        StandardCacheAI,
        StandardCacheSim,
        StandardDuplicateBatch,
        StandardJghPdf,
    )

    try:
        record = await StandardDuplicateBatch.filter(id=batch_id).first()
        if not record:
            return Fail(code="4004", msg="批次记录不存在")

        # ── 参数预处理 ──
        filter_tag_list = [
            t.strip() for t in (filter_tags or "").split(",")
            if t.strip() and t.strip() in _VALID_TAGS
        ]
        hidden_tag_list = [
            t.strip() for t in (hidden_tags or "").split(",")
            if t.strip() and t.strip() in _VALID_TAGS
        ]
        _domain_key = (filter_std_domain or "").strip()
        cross_domain = _CROSS_DOMAIN.get(_domain_key)  # (source_domain, target_domain) or None
        simple_domain = _domain_key if not cross_domain and _domain_key in ("国家标准", "行业标准") else None

        # ── 构造 SQL WHERE 条件 ──
        where = ["n.batch_id = %s"]
        params: list = [batch_id]

        if only_found:
            where.append("n.found = 1")
        if only_need_attention:
            where.append("n.need_attention = 1")
        if filter_standard_no and filter_standard_no.strip():
            where.append("n.standard_no LIKE %s")
            params.append(f"%{_escape_like(filter_standard_no.strip())}%")
        if filter_standard_name and filter_standard_name.strip():
            where.append("n.standard_name LIKE %s")
            params.append(f"%{_escape_like(filter_standard_name.strip())}%")

        # main_has_res：仅约束主标准本身在 jgh_pdf 中是否存在（SQL 部分）。
        # 完整语义"主在 + 相似至少一个在"在拿到当前页后再过滤。
        if main_has_res_filter == 'has_res':
            where.append(
                "EXISTS (SELECT 1 FROM standard_jgh_pdf p "
                "WHERE p.standard_no = n.standard_no)"
            )
        elif main_has_res_filter == 'no_res':
            where.append(
                "NOT EXISTS (SELECT 1 FROM standard_jgh_pdf p "
                "WHERE p.standard_no = n.standard_no)"
            )

        # 领域筛选（SQL 层只过滤源标准 domain；目标 domain + need_attention 在 Python 后置处理）
        if cross_domain:
            where.append(
                "EXISTS (SELECT 1 FROM standard_base_info b "
                "WHERE b.standard_no = n.standard_no AND b.std_domain = %s)"
            )
            params.append(cross_domain[0])
        elif simple_domain:
            where.append(
                "EXISTS (SELECT 1 FROM standard_base_info b "
                "WHERE b.standard_no = n.standard_no AND b.std_domain = %s)"
            )
            params.append(simple_domain)

        # filter_tags：JSON_TABLE 拍平 similar_standards 后判断 tag bool
        # mode='all' 旧语义是"主记录所有 similar 合并的 tag 集合包含每个 tag" → 每个 tag 一段独立 EXISTS
        # mode='any' → 一段 EXISTS 内 OR
        def _tag_path_clause(tag: str) -> str:
            # 路径里的引号需转义
            return f"t INT PATH '$.tags.\"{tag}\"' DEFAULT '0' ON EMPTY"

        if filter_tag_list:
            if filter_tags_mode == 'all':
                for tag in filter_tag_list:
                    where.append(
                        "EXISTS (SELECT 1 FROM JSON_TABLE(n.similar_standards, "
                        f"'$[*]' COLUMNS ({_tag_path_clause(tag)})) jt WHERE jt.t = 1)"
                    )
            else:
                cols = ", ".join(
                    f"t{i} INT PATH '$.tags.\"{t}\"' DEFAULT '0' ON EMPTY"
                    for i, t in enumerate(filter_tag_list)
                )
                or_expr = " OR ".join(f"jt.t{i} = 1" for i in range(len(filter_tag_list)))
                where.append(
                    f"EXISTS (SELECT 1 FROM JSON_TABLE(n.similar_standards, "
                    f"'$[*]' COLUMNS ({cols})) jt WHERE {or_expr})"
                )

        if hidden_tag_list:
            cols = ", ".join(
                f"t{i} INT PATH '$.tags.\"{t}\"' DEFAULT '0' ON EMPTY"
                for i, t in enumerate(hidden_tag_list)
            )
            or_expr = " OR ".join(f"jt.t{i} = 1" for i in range(len(hidden_tag_list)))
            where.append(
                f"NOT EXISTS (SELECT 1 FROM JSON_TABLE(n.similar_standards, "
                f"'$[*]' COLUMNS ({cols})) jt WHERE {or_expr})"
            )

        where_sql = " AND ".join(where)

        # ── 排序 ──
        sort_map = {
            "standard_no": "n.standard_no",
            "standard_name": "n.standard_name",
            "need_attention": "(CASE WHEN n.need_attention=1 THEN 2 WHEN n.found=1 THEN 1 ELSE 0 END)",
            "similar_count": "COALESCE(JSON_LENGTH(n.similar_standards), 0)",
        }
        order_sql = "n.id ASC"
        if sort_field in sort_map and sort_order in ("ascend", "descend"):
            direction = "DESC" if sort_order == "descend" else "ASC"
            order_sql = f"{sort_map[sort_field]} {direction}, n.id ASC"

        # ── 是否需要"主页面后置过滤"（拉两页 buffer 以缓解 main_has_res / 跨域完整语义）──
        need_post_filter = main_has_res_filter in ('has_res', 'no_res') or bool(cross_domain)

        offset = max(0, (current - 1) * size)
        # 后置过滤可能筛掉一些行，因此适度放大 LIMIT。完整语义靠 _post_filter_main_has_res
        # 实现，结果总数 total 会按 SQL 层 COUNT 估算（保持旧行为：旧实现里 total 也是 SQL 后置算）。
        sql_limit = size if not need_post_filter else max(size * 4, 40)

        conn = Tortoise.get_connection("conn_standard")

        data_sql = (
            "SELECT n.id, n.standard_no, n.standard_name, n.use_range, n.found, "
            "n.error, n.need_attention, n.general_evaluation, n.suggestion, "
            "n.similar_standards "
            f"FROM standard_duplicate_name n WHERE {where_sql} "
            f"ORDER BY {order_sql} LIMIT %s OFFSET %s"
        )
        # 跨域模式下，COUNT 也要加目标 domain + need_attention 的 EXISTS 子查询，
        # 保证 pagination.total 与 Python 后置过滤结果一致。
        if cross_domain:
            cross_count_extra = (
                " AND EXISTS ("
                "SELECT 1 FROM JSON_TABLE(n.similar_standards, '$[*]' COLUMNS ("
                "  sno VARCHAR(200) PATH '$.standard_no',"
                "  na  INT         PATH '$.need_attention' DEFAULT '0' ON EMPTY"
                ")) jt2 JOIN standard_base_info b2 ON b2.standard_no = jt2.sno"
                " WHERE b2.std_domain = %s AND jt2.na = 1)"
            )
            count_sql = (
                f"SELECT COUNT(*) AS c FROM standard_duplicate_name n WHERE {where_sql}{cross_count_extra}"
            )
            count_params = params + [cross_domain[1]]
        else:
            count_sql = (
                f"SELECT COUNT(*) AS c FROM standard_duplicate_name n WHERE {where_sql}"
            )
            count_params = params

        # 并发拉数据 + 总数
        import asyncio
        data_task = conn.execute_query(data_sql, params + [sql_limit, offset])
        count_task = conn.execute_query(count_sql, count_params)
        (_, data_rows), (_, count_rows) = await asyncio.gather(data_task, count_task)
        total = int(count_rows[0]["c"]) if count_rows else 0

        # ── 反序列化行 ──
        import json
        items: List[dict] = []
        for row in data_rows:
            sim = row.get("similar_standards")
            if isinstance(sim, str):
                try:
                    sim = json.loads(sim) if sim else []
                except Exception:
                    sim = []
            elif sim is None:
                sim = []
            items.append({
                "id": row["id"],
                "standard_no": row["standard_no"],
                "standard_name": row["standard_name"],
                "use_range": row["use_range"],
                "found": bool(row["found"]),
                "error": row["error"],
                "need_attention": bool(row["need_attention"]),
                "general_evaluation": row["general_evaluation"],
                "suggestion": row["suggestion"],
                "similar_standards": sim or [],
            })

        # ── 当前页范围内的小批量 IN 查询：注入动态字段 ──
        page_main_nos = {it["standard_no"] for it in items if it["standard_no"]}
        page_similar_nos = {
            s.get("standard_no")
            for it in items if it["found"] and it["similar_standards"]
            for s in it["similar_standards"] if s.get("standard_no")
        }
        all_check_nos = page_main_nos | page_similar_nos

        existing_pdf_nos: set = set()
        if all_check_nos:
            pdf_records = await StandardJghPdf.filter(
                standard_no__in=list(all_check_nos)
            ).values_list("standard_no", flat=True)
            existing_pdf_nos = set(pdf_records)

        base_info_id_map: dict = {}
        similar_domain_map: dict = {}
        if page_similar_nos:
            base_info_records = await StandardBaseInfo.filter(
                standard_no__in=list(page_similar_nos)
            ).values_list("standard_no", "id", "std_domain")
            base_info_id_map = {no: str(bid) for no, bid, _ in base_info_records}
            similar_domain_map = {no: domain for no, _, domain in base_info_records}

        page_found_main_nos = [it["standard_no"] for it in items if it["found"]]
        ai_pairs: set = set()
        sim_pairs: set = set()
        if page_similar_nos and page_found_main_nos:
            ai_fwd = await StandardCacheAI.filter(
                source_standard_no__in=page_found_main_nos,
                target_standard_no__in=list(page_similar_nos),
                is_valid=True,
            ).values_list("source_standard_no", "target_standard_no")
            ai_rev = await StandardCacheAI.filter(
                source_standard_no__in=list(page_similar_nos),
                target_standard_no__in=page_found_main_nos,
                is_valid=True,
            ).values_list("source_standard_no", "target_standard_no")
            ai_pairs = {(s, t) for s, t in ai_fwd} | {(t, s) for s, t in ai_rev}

            sim_fwd = await StandardCacheSim.filter(
                source_standard_no__in=page_found_main_nos,
                target_standard_no__in=list(page_similar_nos),
                is_valid=True,
            ).values_list("source_standard_no", "target_standard_no")
            sim_rev = await StandardCacheSim.filter(
                source_standard_no__in=list(page_similar_nos),
                target_standard_no__in=page_found_main_nos,
                is_valid=True,
            ).values_list("source_standard_no", "target_standard_no")
            sim_pairs = {(s, t) for s, t in sim_fwd} | {(t, s) for s, t in sim_rev}

        # 注入动态字段
        for it in items:
            if not (it["found"] and it["similar_standards"]):
                continue
            for s in it["similar_standards"]:
                if "需要关注" in s:
                    s["need_attention"] = s.pop("需要关注")
                std_no = s.get("standard_no")
                s["has_res"] = std_no in existing_pdf_nos if std_no else False
                if std_no and std_no in base_info_id_map:
                    s["id"] = base_info_id_map[std_no]
                s["has_ai_comparison_cache"] = (it["standard_no"], std_no) in ai_pairs if std_no else False
                s["has_sim_cache"] = (it["standard_no"], std_no) in sim_pairs if std_no else False
                s["std_domain"] = similar_domain_map.get(std_no) if std_no else None

        # ── 跨域后置过滤 ──
        # 条件：similar_standards 里至少有一条 target_domain + need_attention=True
        # 同时把 similar_standards 裁剪为仅保留 target_domain 的条目
        if cross_domain:
            target_domain = cross_domain[1]
            filtered_items = []
            for it in items:
                if it["found"] and it["similar_standards"]:
                    target_similars = [
                        s for s in it["similar_standards"]
                        if s.get("std_domain") == target_domain
                    ]
                    if any(s.get("need_attention") for s in target_similars):
                        it["similar_standards"] = target_similars
                        filtered_items.append(it)
            items = filtered_items

        # ── main_has_res 完整语义后置过滤 ──
        # SQL 层只保证主标准在 jgh_pdf；这里继续要求"相似中至少一个有 PDF"
        if main_has_res_filter == 'has_res':
            items = [
                it for it in items
                if it["similar_standards"]
                and any(s.get("has_res") for s in it["similar_standards"])
            ]
        elif main_has_res_filter == 'no_res':
            items = [
                it for it in items
                if (not it["similar_standards"])
                or all(not s.get("has_res") for s in it["similar_standards"])
            ]

        # ── has_res_filter：剔除 similar_standards 数组内元素（不剔主记录）──
        if has_res_filter == 'has_res':
            for it in items:
                if it["found"] and it["similar_standards"]:
                    it["similar_standards"] = [s for s in it["similar_standards"] if s.get("has_res")]
        elif has_res_filter == 'no_res':
            for it in items:
                if it["found"] and it["similar_standards"]:
                    it["similar_standards"] = [s for s in it["similar_standards"] if not s.get("has_res")]

        # 后置过滤可能让 items > size，截到 size
        if need_post_filter:
            items = items[:size]

        return Success(
            data={
                "record": {
                    "id": record.id,
                    "batch_name": record.batch_name,
                    "create_time": record.create_time.isoformat() if record.create_time else None,
                    "total_count": record.total_count,
                    "success_count": record.success_count,
                    "failed_count": record.failed_count,
                    "duplicate_count": record.duplicate_count,
                    "status": record.status,
                    "results": items,
                    "remark": record.remark,
                },
                "pagination": {
                    "current": current,
                    "size": size,
                    "total": total,
                    "pages": (total + size - 1) // size if total > 0 else 0,
                }
            },
            msg="获取详情成功"
        )
    except Exception as e:
        logger.exception(f"获取批次详情失败 batch_id={batch_id}")
        return Fail(code="5000", msg=f"获取详情失败: {str(e)}")


@router.get("/batch-stats/{batch_id}", summary="批次标签统计分析")
async def get_batch_stats(batch_id: int, filter_std_domain: str = None):
    """获取批次的标签统计数据（用于可视化图表）。

    用单条 SQL 聚合代替全量加载：JSON_TABLE 把每条主记录的 similar
    数组拍平，按 n.id 分组取每个 tag 的 MAX 得到主记录级 bool，
    外层再 SUM 得到 10 个统计值。支持 filter_std_domain 筛选（含跨域模式）。
    """
    from tortoise import Tortoise

    from app.models.standard import StandardDuplicateBatch

    try:
        record = await StandardDuplicateBatch.filter(id=batch_id).first()
        if not record:
            return Fail(code="4004", msg="批次记录不存在")

        _sd_key = (filter_std_domain or "").strip()
        cross_domain_stats = _CROSS_DOMAIN.get(_sd_key)
        simple_domain_stats = _sd_key if not cross_domain_stats and _sd_key in ("国家标准", "行业标准") else None

        extra_where = ""
        sql_params: list = [batch_id]

        if cross_domain_stats:
            # 源标准 domain 过滤
            extra_where += (
                " AND EXISTS (SELECT 1 FROM standard_base_info b "
                "WHERE b.standard_no = n.standard_no AND b.std_domain = %s)"
            )
            sql_params.append(cross_domain_stats[0])
            # 目标 domain + need_attention 过滤（JSON_TABLE 展开相似标准后 JOIN standard_base_info）
            extra_where += (
                " AND EXISTS ("
                "SELECT 1 FROM JSON_TABLE(n.similar_standards, '$[*]' COLUMNS ("
                "  sno VARCHAR(200) PATH '$.standard_no',"
                "  na  INT         PATH '$.need_attention' DEFAULT '0' ON EMPTY"
                ")) jt2 JOIN standard_base_info b2 ON b2.standard_no = jt2.sno"
                " WHERE b2.std_domain = %s AND jt2.na = 1)"
            )
            sql_params.append(cross_domain_stats[1])
        elif simple_domain_stats:
            extra_where += (
                " AND EXISTS (SELECT 1 FROM standard_base_info b "
                "WHERE b.standard_no = n.standard_no AND b.std_domain = %s)"
            )
            sql_params.append(simple_domain_stats)

        sql = f"""
        SELECT
          COUNT(*)                                          AS total,
          SUM(need_attention)                               AS need_attention,
          SUM(has_series)                                   AS cnt_series,
          SUM(has_obj)                                      AS cnt_obj,
          SUM(has_gs)                                       AS cnt_gs,
          SUM(has_overlap)                                  AS cnt_overlap,
          SUM(has_obj=1 OR has_overlap=1)                   AS high_risk,
          SUM(has_obj+has_gs+has_overlap=1)                 AS comb_single,
          SUM(has_obj+has_gs+has_overlap=2)                 AS comb_double,
          SUM(has_obj+has_gs+has_overlap>=3)                AS comb_triple
        FROM (
          SELECT n.id,
            MAX(n.need_attention=1)                                                   AS need_attention,
            MAX(IFNULL(jt.t1, 0))                                                     AS has_series,
            MAX(IFNULL(jt.t2, 0))                                                     AS has_obj,
            MAX(IFNULL(jt.t3, 0))                                                     AS has_gs,
            MAX(IFNULL(jt.t4, 0))                                                     AS has_overlap
          FROM standard_duplicate_name n
          LEFT JOIN JSON_TABLE(n.similar_standards, '$[*]' COLUMNS (
            t1 INT PATH '$.tags."同系列标准"' DEFAULT '0' ON EMPTY,
            t2 INT PATH '$.tags."标准化对象一致"' DEFAULT '0' ON EMPTY,
            t3 INT PATH '$.tags."通用专用关系"' DEFAULT '0' ON EMPTY,
            t4 INT PATH '$.tags."适用范围重叠"' DEFAULT '0' ON EMPTY
          )) jt ON 1=1
          WHERE n.batch_id = %s AND n.found = 1{extra_where}
          GROUP BY n.id
        ) m
        """
        conn = Tortoise.get_connection("conn_standard")
        _, rows = await conn.execute_query(sql, sql_params)
        row = rows[0] if rows else {}

        def _i(v):
            return int(v) if v is not None else 0

        total_found = _i(row.get("total"))

        if total_found == 0:
            return Success(
                data={
                    "high_risk_stats": {"need_attention_count": 0, "total_count": 0, "high_risk_tag_count": 0},
                    "tag_ranking": [],
                    "tag_combinations": {"single_tag": 0, "double_tag": 0, "triple_plus": 0},
                    "tag_distribution": {},
                },
                msg="无有效数据"
            )

        tag_counts = {
            "同系列标准": _i(row.get("cnt_series")),
            "标准化对象一致": _i(row.get("cnt_obj")),
            "通用专用关系": _i(row.get("cnt_gs")),
            "适用范围重叠": _i(row.get("cnt_overlap")),
        }
        tag_ranking = [
            {"tag": k, "count": v, "percentage": round(v / total_found * 100, 2)}
            for k, v in sorted(tag_counts.items(), key=lambda x: x[1], reverse=True)
        ]
        tag_distribution = {k: round(v / total_found, 4) for k, v in tag_counts.items()}

        return Success(
            data={
                "high_risk_stats": {
                    "need_attention_count": _i(row.get("need_attention")),
                    "total_count": total_found,
                    "high_risk_tag_count": _i(row.get("high_risk")),
                },
                "tag_ranking": tag_ranking,
                "tag_combinations": {
                    "single_tag": _i(row.get("comb_single")),
                    "double_tag": _i(row.get("comb_double")),
                    "triple_plus": _i(row.get("comb_triple")),
                },
                "tag_distribution": tag_distribution,
            },
            msg="统计成功"
        )
    except Exception as e:
        logger.exception(f"批次统计失败 batch_id={batch_id}")
        return Fail(code="5000", msg=f"统计失败: {str(e)}")


@router.get("/batch-export/{batch_id}", summary="导出批次详情到 Excel")
async def export_batch_to_excel(batch_id: int):
    """导出批次详情到 Excel。

    并行跑两条 SQL：
    1) 主表行（不含 similar_standards JSON 列）
    2) 相似明细 — 用 JSON_TABLE 在 OceanBase 端把数组拍平为关系行，
       避免传输 LongText JSON。
    再用 openpyxl write_only 模式流式写盘，30 万行规模也不会爆内存。
    """
    import asyncio
    import time

    from tortoise import Tortoise

    from app.models.standard import StandardDuplicateBatch

    t_start = time.time()
    try:
        record = await StandardDuplicateBatch.filter(id=batch_id).first()
        if not record:
            raise HTTPException(status_code=404, detail="批次记录不存在")

        logger.info(f"[Export] 开始导出 batch_id={batch_id} batch_name={record.batch_name!r}")

        conn = Tortoise.get_connection("conn_standard")

        # 主表（无 JSON 列），同时用 JSON_LENGTH 在 SQL 端算出 similar_count
        detail_sql = """
        SELECT id, standard_no, standard_name, use_range, found, error,
               need_attention, general_evaluation, suggestion,
               COALESCE(JSON_LENGTH(similar_standards), 0) AS similar_count
        FROM standard_duplicate_name
        WHERE batch_id = %s
        ORDER BY id
        """

        # 相似明细：JSON_TABLE 拍平 + 主表的 standard_no/name/use_range/need_attention 一并 SELECT
        similar_sql = """
        SELECT
          n.standard_no AS main_standard_no,
          n.standard_name AS main_standard_name,
          n.use_range AS main_use_range,
          n.need_attention AS main_need_attention,
          jt.sn, jt.cn, jt.ur, jt.ls,
          jt.t1, jt.t2, jt.t3, jt.t4, jt.rd
        FROM standard_duplicate_name n
        JOIN JSON_TABLE(n.similar_standards, '$[*]' COLUMNS (
          sn VARCHAR(200) PATH '$.standard_no',
          cn TEXT PATH '$.cname',
          ur TEXT PATH '$.use_range',
          ls DOUBLE PATH '$.llm_score' DEFAULT '0' ON EMPTY,
          t1 INT PATH '$.tags."同系列标准"' DEFAULT '0' ON EMPTY,
          t2 INT PATH '$.tags."标准化对象一致"' DEFAULT '0' ON EMPTY,
          t3 INT PATH '$.tags."通用专用关系"' DEFAULT '0' ON EMPTY,
          t4 INT PATH '$.tags."适用范围重叠"' DEFAULT '0' ON EMPTY,
          rd TEXT PATH '$.relation_desc'
        )) jt ON 1=1
        WHERE n.batch_id = %s AND n.found = 1
        ORDER BY n.id
        """

        async def _fetch_detail():
            t0 = time.time()
            _, rows = await conn.execute_query(detail_sql, [batch_id])
            logger.info(f"[Export] batch_id={batch_id} 主表查询完成 rows={len(rows)} elapsed={time.time()-t0:.1f}s")
            return rows

        async def _fetch_similar():
            t0 = time.time()
            _, rows = await conn.execute_query(similar_sql, [batch_id])
            logger.info(f"[Export] batch_id={batch_id} 相似明细 JSON_TABLE 拍平完成 rows={len(rows)} elapsed={time.time()-t0:.1f}s")
            return rows

        # 并行跑两条 SQL：耗时 = max 而非 sum
        detail_rows, similar_rows = await asyncio.gather(_fetch_detail(), _fetch_similar())

        # openpyxl 是 CPU 密集，放到线程避免阻塞 event loop
        t_excel = time.time()
        logger.info(f"[Export] batch_id={batch_id} 开始生成 Excel detail={len(detail_rows)} similar={len(similar_rows)}")
        output = await asyncio.to_thread(
            _generate_batch_excel, record, list(detail_rows), list(similar_rows)
        )
        size_mb = output.getbuffer().nbytes / 1024 / 1024
        logger.info(
            f"[Export] batch_id={batch_id} Excel 生成完成 size={size_mb:.2f}MB "
            f"excel_elapsed={time.time()-t_excel:.1f}s total_elapsed={time.time()-t_start:.1f}s"
        )

        filename = f"批次查重结果_{batch_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        return StreamingResponse(
            output,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={
                "Content-Disposition": f"attachment; filename*=UTF-8''{quote(filename)}",
                "Access-Control-Expose-Headers": "Content-Disposition"
            }
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.exception(f"导出Excel失败 batch_id={batch_id} elapsed={time.time()-t_start:.1f}s")
        raise HTTPException(status_code=500, detail=f"导出失败: {str(e)}")


@router.get("/export-standard-word", summary="导出单个标准的比对结果为Word")
async def export_standard_to_word(batch_id: int, standard_no: str):
    """导出单个标准的比对结果为 Word 文档"""
    from app.models.standard import StandardDuplicateBatch, StandardDuplicateName

    try:
        logger.info(f"开始导出Word: batch_id={batch_id}, standard_no={standard_no}")

        record = await StandardDuplicateBatch.filter(id=batch_id).first()
        if not record:
            raise HTTPException(status_code=404, detail="批次记录不存在")

        item = await StandardDuplicateName.filter(batch_id=batch_id, standard_no=standard_no).first()
        if not item:
            raise HTTPException(status_code=404, detail="未找到该标准的比对结果")
        if not item.found:
            raise HTTPException(status_code=400, detail="该标准未找到或分析失败")

        standard_result = {
            "standard_no": item.standard_no,
            "standard_name": item.standard_name,
            "use_range": item.use_range,
            "found": item.found,
            "need_attention": item.need_attention,
            "general_evaluation": item.general_evaluation,
            "suggestion": item.suggestion,
            "similar_standards": item.similar_standards or []
        }
        output = await _build_word_document(standard_result, standard_no)
        filename = f"{standard_no}_比对结果_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        return StreamingResponse(
            output,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename*=UTF-8''{quote(filename)}",
                "Access-Control-Expose-Headers": "Content-Disposition"
            }
        )
    except HTTPException:
        raise
    except Exception as e:
        import traceback
        logger.error(f"导出Word失败: {e}")
        logger.error(f"完整traceback:\n{traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"导出失败: {str(e)}")


@router.post("/batch-resume/{batch_id}", summary="续跑未完成批次")
async def resume_batch(batch_id: int):
    """
    手动续跑某个 processing 批次中剩余的标准（pending/running/failed）。
    服务重启后已由启动钩子自动调度，此接口提供前端手动触发能力。
    """
    import asyncio as _asyncio

    from app.langchain.agents.tasks.standard.standard_deduplication_agent import resume_batch_dispatch
    from app.models.standard import StandardDuplicateBatch

    record = await StandardDuplicateBatch.filter(id=batch_id).first()
    if not record:
        return Fail(code="4004", msg="批次记录不存在")

    # 后台触发，不阻塞 HTTP 请求
    _asyncio.create_task(resume_batch_dispatch(batch_id))
    return Success(data=None, msg="已触发续跑，请稍后刷新查看进度")


class BatchMergeRequest(BaseModel):
    source_batch_ids: List[int] = Field(..., description="待合并的批次ID列表", min_length=2)
    batch_name: Optional[str] = Field(None, description="合并后的新批次名称", max_length=200)
    remark: Optional[str] = Field(None, description="备注")


@router.post("/batch-merge", summary="合并多个批次")
async def merge_batches(body: BatchMergeRequest):
    """
    合并多个批次为一个新批次。

    规则：
    - 所有源批次的 pool_id 必须一致，否则拒绝
    - 同一 standard_no 在多个源批次中存在时，按 update_time 最新一条保留
    - 原批次保留，仅复制 name 记录到新批次
    """
    from app.models.standard import StandardDuplicateBatch, StandardDuplicateName

    source_ids = list(dict.fromkeys(body.source_batch_ids))  # 去重保序
    if len(source_ids) < 2:
        return Fail(code="4000", msg="至少需要选择 2 个批次")

    sources = await StandardDuplicateBatch.filter(id__in=source_ids).all()
    if len(sources) != len(source_ids):
        found_ids = {s.id for s in sources}
        missing = [i for i in source_ids if i not in found_ids]
        return Fail(code="4004", msg=f"批次不存在：{missing}")

    # pool_id 必须一致
    pool_ids = {s.pool_id for s in sources}
    if len(pool_ids) > 1:
        return Fail(code="4000", msg="所选批次的查重池不一致，无法合并")
    pool_id = pool_ids.pop()

    # 不允许包含 processing 中的批次（避免合并到一半的数据）
    processing = [s.id for s in sources if s.status == "processing"]
    if processing:
        return Fail(code="4000", msg=f"以下批次仍在处理中，无法合并：{processing}")

    # 拉取所有源批次的 name 记录，按 update_time 倒序，同标准号取首条
    all_items = await StandardDuplicateName.filter(
        batch_id__in=source_ids
    ).order_by("-update_time", "-id").all()

    picked: dict[str, StandardDuplicateName] = {}
    for item in all_items:
        if item.standard_no not in picked:
            picked[item.standard_no] = item

    if not picked:
        return Fail(code="4000", msg="选中的批次没有可合并的记录")

    # 创建新批次
    success_count = sum(1 for v in picked.values() if v.found)
    failed_count = len(picked) - success_count
    duplicate_count = sum(1 for v in picked.values() if v.need_attention)

    new_batch = await StandardDuplicateBatch.create(
        batch_name=body.batch_name or f"合并批次（共 {len(picked)} 个标准）",
        pool_id=pool_id,
        total_count=len(picked),
        success_count=success_count,
        failed_count=failed_count,
        duplicate_count=duplicate_count,
        status="completed",
        remark=body.remark or f"由批次 {source_ids} 合并而来",
    )

    # 复制 name 记录到新批次
    new_records = [
        StandardDuplicateName(
            batch_id=new_batch.id,
            standard_no=item.standard_no,
            standard_name=item.standard_name,
            use_range=item.use_range,
            found=item.found,
            error=item.error,
            need_attention=item.need_attention,
            general_evaluation=item.general_evaluation,
            suggestion=item.suggestion,
            similar_standards=item.similar_standards,
            task_status=item.task_status,
        )
        for item in picked.values()
    ]
    await StandardDuplicateName.bulk_create(new_records, batch_size=200)

    logger.info(
        f"[BatchMerge] 合并完成 source={source_ids} → new_batch={new_batch.id}，"
        f"共 {len(picked)} 条标准"
    )

    return Success(
        data={
            "batch_id": new_batch.id,
            "total": len(picked),
            "source_batch_ids": source_ids,
        },
        msg=f"合并成功，新批次包含 {len(picked)} 个标准"
    )


@router.delete("/batch-delete/{batch_id}", summary="删除批次记录")
async def delete_batch_record(batch_id: int):
    """删除指定批次的查重记录"""
    from app.models.standard import StandardDuplicateBatch

    try:
        record = await StandardDuplicateBatch.filter(id=batch_id).first()
        if not record:
            return Fail(code="4004", msg="批次记录不存在")
        await record.delete()
        return Success(data=None, msg="删除成功")
    except Exception as e:
        logger.error(f"删除批次记录失败: {e}")
        return Fail(code="5000", msg=f"删除失败: {str(e)}")


class UpdateSimilarAttentionRequest(BaseModel):
    standard_no: str = Field(..., description="主标准编号")
    similar_standard_no: str = Field(..., description="相似标准编号")
    need_attention: bool = Field(..., description="是否需要关注")


@router.patch("/batch/{batch_id}/similar-attention", summary="更新相似标准的需要关注标记")
async def update_similar_attention(batch_id: int, body: UpdateSimilarAttentionRequest):
    """手动标记某条相似标准是否需要关注，并同步更新主标准的需要关注字段"""
    from app.models.standard import StandardDuplicateName

    try:
        item = await StandardDuplicateName.filter(
            batch_id=batch_id,
            standard_no=body.standard_no
        ).first()
        if not item:
            return Fail(code="4004", msg="未找到对应的标准记录")

        similars = item.similar_standards or []
        updated = False

        for s in similars:
            if s.get("standard_no") == body.similar_standard_no:
                s["need_attention"] = body.need_attention
                updated = True

        if not updated:
            return Fail(code="4004", msg="未找到对应的相似标准条目")

        item.similar_standards = similars
        item.need_attention = any(s.get("need_attention", False) for s in similars)
        await item.save(update_fields=["similar_standards", "need_attention"])
        return Success(data=None, msg="更新成功")
    except Exception as e:
        logger.error(f"更新需要关注失败: {e}")
        return Fail(code="5000", msg=f"更新失败: {str(e)}")


@router.get("/similar-standards-by-no", summary="根据标准号或记录ID获取相似标准列表（支持筛选/排序/分页）")
async def get_similar_standards_by_no(
    standard_no: str = None,
    record_id: int = None,
    # 筛选
    filter_tags: str = None,          # 显示筛选：逗号分隔的标签，如 "标准化对象一致,通用专用关系"
    hidden_tags: str = None,          # 隐藏筛选：逗号分隔的标签
    filter_tags_mode: str = 'any',    # any / all
    has_res_filter: str = None,       # has_res / no_res / 空表示不限
    llm_score_min: float = None,      # llm_score 最低分
    need_attention_only: bool = False, # 仅需要关注
    filter_std_domain: str = None,    # 国家标准 / 行业标准 / 国家标准对行业标准 / 行业标准对国家标准
    # 排序
    sort_by: str = 'llm_score',       # llm_score / tag
    # 分页
    page: int = 1,
    page_size: int = 10,
):
    """
    根据标准编号或记录ID查询相似标准列表。
    record_id 优先；若只传 standard_no，则取该标准编号最新的查重记录。
    筛选/排序/分页均在服务端完成，tags 字段统一转换为数组格式。
    """
    from app.models.standard import StandardDuplicateName, StandardJghPdf

    _TAG_ALL = {'同系列标准', '标准化对象一致', '通用专用关系', '适用范围重叠'}
    _TAG_PRIORITY = {'标准化对象一致': 1, '适用范围重叠': 2, '通用专用关系': 3, '同系列标准': 4}

    def _to_tag_list(raw) -> list:
        """将 tags 字段统一转为数组（兼容旧 dict 和新 list 格式）"""
        if isinstance(raw, list):
            return [t for t in raw if t in _TAG_ALL]
        if isinstance(raw, dict):
            return [k for k, v in raw.items() if v and k in _TAG_ALL]
        return []

    try:
        if record_id is not None:
            item = await StandardDuplicateName.filter(id=record_id).first()
            if not item:
                return Fail(code="4004", msg="未找到对应的查重记录")
        elif standard_no:
            item = await StandardDuplicateName.filter(
                standard_no=standard_no
            ).order_by("-id").first()
            if not item:
                return Fail(code="4004", msg="未找到该标准的查重记录")
        else:
            return Fail(code="4000", msg="请传入 record_id 或 standard_no")

        all_similars = item.similar_standards or []

        # 动态注入 has_res + std_domain（DB 中该字段为默认值，需实时查询）
        from app.models.standard import StandardBaseInfo
        similar_nos = [s.get('standard_no') for s in all_similars if s.get('standard_no')]
        if similar_nos:
            pdf_nos = set(await StandardJghPdf.filter(
                standard_no__in=similar_nos
            ).values_list('standard_no', flat=True))
            domain_records = await StandardBaseInfo.filter(
                standard_no__in=similar_nos
            ).values_list('standard_no', 'std_domain')
            domain_map = {no: domain for no, domain in domain_records}
        else:
            pdf_nos = set()
            domain_map = {}

        # 动态注入全文相似度（来自 standard_cache_sim 缓存，双向查询取有效记录）
        from app.models.standard.cache_sim import StandardCacheSim
        source_no = item.standard_no
        full_sim_map = {}
        if source_no and similar_nos:
            # 正向：source=当前标准, target=相似标准
            fwd = await StandardCacheSim.filter(
                source_standard_no=source_no,
                target_standard_no__in=similar_nos,
                is_valid=True
            ).values_list('target_standard_no', 'similarity_percentage')
            for t_no, pct in fwd:
                full_sim_map[t_no] = pct
            # 反向：source=相似标准, target=当前标准
            rev = await StandardCacheSim.filter(
                source_standard_no__in=similar_nos,
                target_standard_no=source_no,
                is_valid=True
            ).values_list('source_standard_no', 'similarity_percentage')
            for s_no, pct in rev:
                full_sim_map.setdefault(s_no, pct)

        # 解析跨域筛选
        _sd_key = (filter_std_domain or "").strip()
        cross_domain_detail = _CROSS_DOMAIN.get(_sd_key)

        # tags 转数组 + 注入 has_res / std_domain（操作副本，不修改 DB 数据）
        import copy
        all_similars = copy.deepcopy(all_similars)
        for s in all_similars:
            # 兼容旧数据中的中文字段名
            if '需要关注' in s:
                s['need_attention'] = s.pop('需要关注')
            s['tags'] = _to_tag_list(s.get('tags'))
            s['has_res'] = s.get('standard_no') in pdf_nos
            s['std_domain'] = domain_map.get(s.get('standard_no'))
            s['full_text_similarity'] = full_sim_map.get(s.get('standard_no'))

        # 跨域筛选：仅保留对侧 domain 的相似标准（统计同步裁剪）
        if cross_domain_detail:
            target_domain_detail = cross_domain_detail[1]
            all_similars = [s for s in all_similars if s.get('std_domain') == target_domain_detail]

        # ── 统计（基于领域筛选后的全量）──
        stats = {
            'total': len(all_similars),
            'need_attention': sum(1 for s in all_similars if s.get('need_attention') is True),
            'ignored': sum(1 for s in all_similars if s.get('need_attention') is False),
            'same_series': sum(1 for s in all_similars if '同系列标准' in s['tags']),
        }

        # ── 筛选 ──
        filtered = all_similars

        if filter_tags:
            tag_list = [t.strip() for t in filter_tags.split(',') if t.strip()]
            if tag_list:
                if filter_tags_mode == 'all':
                    filtered = [s for s in filtered if all(t in s['tags'] for t in tag_list)]
                else:
                    filtered = [s for s in filtered if any(t in s['tags'] for t in tag_list)]

        if hidden_tags:
            hide_list = [t.strip() for t in hidden_tags.split(',') if t.strip()]
            if hide_list:
                filtered = [s for s in filtered if not any(t in s['tags'] for t in hide_list)]

        if has_res_filter == 'has_res':
            filtered = [s for s in filtered if s.get('has_res')]
        elif has_res_filter == 'no_res':
            filtered = [s for s in filtered if not s.get('has_res')]

        if llm_score_min is not None:
            filtered = [s for s in filtered if (s.get('llm_score') or 0) >= llm_score_min]

        if need_attention_only:
            filtered = [s for s in filtered if s.get('need_attention') is True]

        # ── 排序 ──
        if sort_by == 'tag':
            def _tag_sort_key(s):
                tags = s['tags']
                if not tags:
                    return (999, 0)
                best = min(_TAG_PRIORITY.get(t, 999) for t in tags)
                return (best, -len(tags))
            filtered.sort(key=_tag_sort_key)
        else:
            filtered.sort(key=lambda s: (s.get('llm_score') or 0), reverse=True)

        # ── 分页 ──
        total = len(filtered)
        start = (page - 1) * page_size
        paged = filtered[start:start + page_size]

        return Success(data={
            'list': paged,
            'total': total,
            'stats': stats,
        })
    except Exception as e:
        logger.error(f"查询相似标准列表失败: {e}")
        return Fail(code="5000", msg=f"查询失败: {str(e)}")


__all__ = ["router"]
